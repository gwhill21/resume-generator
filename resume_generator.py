print("Starting resume generation...")

from faker import Faker
from docx import Document
from docx.shared import Pt
from io import BytesIO
import zipfile
import os
import random
import pandas as pd

# === Load data ===
colleges_df = pd.read_csv("us_colleges.csv")
schools = colleges_df['name'].tolist()

companies_df = pd.read_csv("fortune_500.csv")
companies = companies_df['company'].tolist()

skills_df = pd.read_csv("trending_skills.csv")
all_skills = skills_df["skill"].dropna().tolist()

# === Initialize Faker ===
fake = Faker("en_US")

# === Function profiles ===
functions = {
    "HR": {
        "titles": {
            "Entry": ["HR Assistant", "Recruiting Coordinator"],
            "Mid": ["HR Business Partner", "Talent Acquisition Manager"],
            "Senior": ["Director of People", "VP of Human Resources"]
        },
        "experience": [
            "Coordinated onboarding for 100+ employees annually",
            "Implemented HRIS tools to automate core processes",
            "Reduced time-to-hire by 30% through new sourcing strategies",
            "Managed benefits administration and policy compliance"
        ],
        "skills": ["Workday", "HRIS", "Employee Relations", "Talent Acquisition", "Onboarding"]
    },
    "Finance": {
        "titles": {
            "Entry": ["Financial Analyst I", "Junior Accountant"],
            "Mid": ["FP&A Analyst", "Finance Manager"],
            "Senior": ["Director of Finance", "VP of FP&A"]
        },
        "experience": [
            "Created and maintained monthly financial reports and models",
            "Improved budgeting accuracy by integrating BI tools",
            "Managed forecasts and variance analysis for $10M+ budgets",
            "Led cost-optimization efforts, reducing expenses by 12%"
        ],
        "skills": ["Excel", "SAP", "Hyperion", "Forecasting", "Financial Modeling"]
    },
    "Marketing": {
        "titles": {
            "Entry": ["Marketing Coordinator", "Content Assistant"],
            "Mid": ["Marketing Manager", "Brand Strategist"],
            "Senior": ["Director of Marketing", "VP of Growth"]
        },
        "experience": [
            "Led B2B demand gen campaigns with 300% ROI",
            "Owned SEO/SEM strategy and increased organic traffic",
            "Managed $1M ad spend across digital platforms",
            "Launched product rebrands and go-to-market campaigns"
        ],
        "skills": ["Google Analytics", "HubSpot", "SEO", "Content Strategy", "Paid Media"]
    },
    "Legal": {
        "titles": {
            "Entry": ["Legal Assistant", "Contract Coordinator"],
            "Mid": ["Corporate Counsel", "Compliance Manager"],
            "Senior": ["Head of Legal", "VP of Compliance"]
        },
        "experience": [
            "Drafted and reviewed vendor and SaaS agreements",
            "Provided legal support for mergers and acquisitions",
            "Managed compliance with GDPR, CCPA, and SOX",
            "Led internal investigations and policy development"
        ],
        "skills": ["Legal Research", "Contract Review", "Compliance", "GDPR", "DocuSign"]
    },
    "Engineering": {
        "titles": {
            "Entry": ["Software Engineer I", "QA Tester"],
            "Mid": ["Backend Engineer", "DevOps Engineer"],
            "Senior": ["Engineering Manager", "Director of Engineering"]
        },
        "experience": [
            "Built scalable microservices with REST and GraphQL",
            "Improved CI/CD pipelines and deployment automation",
            "Managed distributed systems using Kubernetes and AWS",
            "Reduced technical debt and led code reviews"
        ],
        "skills": ["Python", "Docker", "Kubernetes", "CI/CD", "AWS"]
    },
    "IT": {
        "titles": {
            "Entry": ["Help Desk Analyst", "IT Support Tech"],
            "Mid": ["IT Systems Analyst", "Network Administrator"],
            "Senior": ["Director of IT", "Infrastructure Lead"]
        },
        "experience": [
            "Provided Tier 1–3 support for global user base",
            "Maintained network infrastructure and VPN access",
            "Led migration from on-prem to Azure AD",
            "Automated device provisioning and patching"
        ],
        "skills": ["ServiceNow", "Active Directory", "Windows Server", "Networking", "ITIL"]
    },
    "Data": {
        "titles": {
            "Entry": ["Data Analyst", "Research Assistant"],
            "Mid": ["Data Scientist", "BI Developer"],
            "Senior": ["Director of Data", "Lead Data Engineer"]
        },
        "experience": [
            "Built dashboards to support executive KPIs",
            "Created ETL pipelines with Airflow and Snowflake",
            "Applied machine learning to detect fraud patterns",
            "Led migration of legacy analytics stack to cloud"
        ],
        "skills": ["SQL", "Tableau", "Python", "Snowflake", "Machine Learning"]
    },
    "Customer Support": {
        "titles": {
            "Entry": ["Support Specialist", "Customer Success Rep"],
            "Mid": ["Customer Success Manager", "Support Team Lead"],
            "Senior": ["Head of Support", "VP of CX"]
        },
        "experience": [
            "Resolved over 500 tickets per quarter with 98% CSAT",
            "Developed onboarding flow for new customers",
            "Managed enterprise accounts post-implementation",
            "Decreased churn through proactive engagement"
        ],
        "skills": ["Zendesk", "CRM", "Intercom", "Retention Strategy", "Customer Advocacy"]
    },
    "Product": {
        "titles": {
            "Entry": ["Product Analyst", "Junior PM"],
            "Mid": ["Product Manager", "Technical PM"],
            "Senior": ["Director of Product", "Head of Product"]
        },
        "experience": [
            "Defined product roadmap aligned with OKRs",
            "Collaborated with engineering to ship MVP in 6 weeks",
            "Ran customer interviews and usability studies",
            "Launched 3 new features leading to 25% retention lift"
        ],
        "skills": ["Agile", "User Research", "Figma", "Jira", "Roadmapping"]
    },
    "Supply Chain": {
        "titles": {
            "Entry": ["Logistics Coordinator", "Inventory Analyst"],
            "Mid": ["Supply Chain Planner", "Procurement Manager"],
            "Senior": ["Director of Supply Chain", "Global Logistics Lead"]
        },
        "experience": [
            "Optimized warehouse flow reducing costs by 15%",
            "Managed supplier relationships across 3 continents",
            "Forecasted demand using time-series models",
            "Led ERP implementation across 4 facilities"
        ],
        "skills": ["Oracle SCM", "SAP MM", "Procurement", "Logistics", "ERP"]
    }
}

certifications_by_function = {
    "HR": ["SHRM-CP", "PHR"],
    "Finance": ["CPA", "CFA Level I"],
    "Product": ["CSPO", "Google UX Cert"],
    "Supply Chain": ["CPIM", "Six Sigma Green Belt"],
    "Marketing": ["Google Ads Cert", "HubSpot Inbound"],
    "Legal": ["JD", "CIPP/E"],
    "Engineering": ["AWS Certified DevOps", "Certified Kubernetes Admin"],
    "IT": ["CompTIA A+", "Microsoft Azure Cert"],
    "Data": ["Google Data Cert", "Azure Data Scientist"],
    "Customer Support": ["HDI Support Cert", "CXPA Certification"]
}

levels = ["Entry", "Mid", "Senior"]

# === Set function counts ===
function_counts = {
    "HR": 2,
    "Finance": 2,
    "Marketing": 2,
    "Legal": 2,
    "IT": 2,
    "Data": 2,
    "Engineering": 2,
    "Product": 2,
    "Customer Support": 2,
    "Supply Chain": 3
}

# === Output setup ===
output_dir = "Generated_Zips"
os.makedirs(output_dir, exist_ok=True)

count = sum(function_counts.values())
zip_filename = os.path.join(output_dir, f"resumes_{count}.zip")

with zipfile.ZipFile(zip_filename, mode='w', compression=zipfile.ZIP_DEFLATED) as zf:
    i = 0
    for function, function_total in function_counts.items():
        for _ in range(function_total):
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)

            level = random.choice(levels)
            profile = functions[function]

            job_title = random.choice(profile["titles"][level])
            experience_samples = random.sample(profile["experience"], 2)
            function_skills = profile["skills"]
            trending_skills = random.sample(all_skills, k=random.randint(3, 7))
            combined_skills = list(set(function_skills + trending_skills))
            selected_skills = random.sample(combined_skills, k=min(10, len(combined_skills)))

            school = random.choice(schools)
            graduation_year = random.randint(2010, 2022)
            gpa = round(random.uniform(3.2, 4.0), 2)

            first_name = fake.first_name()
            last_name = fake.last_name()
            name = f"{first_name} {last_name}"
            email = f"{first_name.lower()}.{last_name.lower()}@example.com"
            city = fake.city()
            state = fake.state_abbr()
            zip_code = fake.zipcode_in_state(state)
            location = f"{city}, {state} {zip_code}"
            phone = fake.phone_number()

            doc.add_heading(name, level=1)
            doc.add_paragraph(f"{email} | {phone} | {location}")

            doc.add_heading("Professional Summary", level=2)
            doc.add_paragraph(
                f"{level}-level {function} professional with {random.choice(['over 5 years', 'extensive', 'proven'])} experience in the tech industry. "
                f"Skilled in {', '.join(selected_skills[:3])}, with a strong record of {random.choice(['enabling innovation', 'building scalable systems', 'cross-functional delivery'])}. "
                f"Recognized for {random.choice(['collaborative mindset', 'attention to detail', 'adaptability', 'leadership'])}."
            )

            doc.add_heading("Work Experience", level=2)
            for j in range(random.randint(2, 3)):
                company = random.choice(companies)
                title = random.choice(profile["titles"][level])
                start_year = graduation_year + j + 1
                end_year = start_year + random.choice([1, 2, 3])
                date_range = f"{random.choice(['Jan', 'May', 'Sep'])} {start_year} – {random.choice(['Feb', 'Jul', 'Dec'])} {end_year}"
                para = doc.add_paragraph()
                run = para.add_run(f"{company} – {title} | {date_range}")
                run.bold = True
                for bullet in random.sample(profile["experience"], 2):
                    doc.add_paragraph(f"• {bullet}")

            doc.add_heading("Education", level=2)
            doc.add_paragraph(f"{school}, Bachelor of Science", style='List Bullet')
            doc.add_paragraph(f"• GPA: {gpa} | Graduated: {graduation_year}")

            if level in ["Mid", "Senior"] and random.random() < 0.4:
                mba_school = random.choice(schools)
                doc.add_paragraph(f"{mba_school}, MBA", style='List Bullet')
                doc.add_paragraph(f"• Graduated: {graduation_year + random.randint(4, 8)}")

            doc.add_heading("Skills", level=2)
            doc.add_paragraph(", ".join(selected_skills), style='List Bullet')

            if level in ["Mid", "Senior"] and random.random() < 0.5:
                certs = random.sample(certifications_by_function.get(function, []), k=min(2, len(certifications_by_function.get(function, []))))
                if certs:
                    doc.add_heading("Certifications", level=2)
                    for cert in certs:
                        doc.add_paragraph(cert, style='List Bullet')

            level_clean = level.lower().replace(" ", "_")
            function_clean = function.lower().replace(" ", "_")
            filename = f"{level_clean}_{function_clean}_resume{i+1}.docx"
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            zf.writestr(filename, buffer.read())

            i += 1

print(f"\n✅ {count} resumes saved to:\n{os.path.abspath(zip_filename)}")
